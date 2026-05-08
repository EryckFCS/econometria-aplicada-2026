
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for edge in ('top', 'start', 'bottom', 'end'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcPr.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcPr.append(element)
            for key, value in edge_data.items():
                element.set(qn('w:{}'.format(key)), str(value))

def create_final_master_table():
    output_path = "/home/erick-fcs/Documentos/universidad/07_Ciclo/septimo_ciclo/applied_econometrics_2026/docs/vaults/u1-aa-01-applied-econometrics/ape2-crisis-iess/reports/Tabla_Master_Final_Audit.docx"
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    
    doc = Document()
    
    # Table Header
    p = doc.add_paragraph()
    run = p.add_run("Table 4")
    run.bold = True
    p.add_run("\nUnit root test results (Levels vs First Differences)")
    
    # Data from Log
    # [Var, ADF_L, PP_L, ZA_L, Break_L, ADF_D, PP_D, ZA_D, Break_D, Verdict]
    data = [
        ["AF", "-1.196", "-0.568", "-5.159**", "2011", "-1.460", "-1.755", "-4.771", "2013", "I(1)"],
        ["FL", "-2.269", "-2.637", "-3.414", "2015", "-3.169**", "-5.451***", "-5.901***", "2012", "I(1)"],
        ["EM", "-2.470", "-4.815***", "-5.363**", "2019", "-3.819***", "-7.445***", "-5.225**", "2010", "I(1)"],
        ["GDP", "-1.604", "-1.599", "-3.660", "2013", "-3.390**", "-4.544***", "-4.490", "2011", "I(1)"],
        ["SBU", "-0.961", "-4.038***", "-15.022***", "2014", "-2.522", "-7.455***", "-11.610***", "2020", "I(1)"]
    ]
    
    table = table = doc.add_table(rows=2 + len(data), cols=10)
    
    # Merge Header Row 1
    table.rows[0].cells[1].merge(table.rows[0].cells[4])
    table.rows[0].cells[1].text = "Test-statistics value at Level"
    table.rows[0].cells[5].merge(table.rows[0].cells[8])
    table.rows[0].cells[5].text = "Test-statistics value at first difference"
    table.rows[0].cells[9].text = "I (d)"
    
    # Sub-headers Row 2
    subs = ["Variable", "(ADF)", "(PP)", "(ZA)", "Break", "(ADF)", "(PP)", "(ZA)", "Break", "Verdict"]
    for i, s in enumerate(subs):
        table.rows[1].cells[i].text = s
        
    # Fill Data
    for r_idx, row_vals in enumerate(data):
        for c_idx, val in enumerate(row_vals):
            table.rows[r_idx + 2].cells[c_idx].text = val

    # Formatting
    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_border(cell, top={}, bottom={}, start={}, end={})
    
    # Top Border (Double)
    for i in range(10):
        set_cell_border(table.rows[0].cells[i], top={"sz": 12, "val": "double", "color": "#000000"})
    
    # Mid Border (Single)
    for i in range(10):
        set_cell_border(table.rows[1].cells[i], bottom={"sz": 6, "val": "single", "color": "#000000"})
        
    # Bottom Border (Single)
    for i in range(10):
        set_cell_border(table.rows[-1].cells[i], bottom={"sz": 12, "val": "single", "color": "#000000"})

    # Left align first column
    for row in table.rows:
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Footer
    doc.add_paragraph()
    note = doc.add_paragraph()
    note.add_run("Note: ***, **, * represent significance at 1%, 5%, and 10% levels respectively.").font.size = Pt(9)
    note.add_run("\nAF: Affiliates, FL: Labor Force, EM: Country Risk, GDP: PIB pc, SBU: Basic Salary.").font.size = Pt(9)

    doc.save(output_path)
    print(f"Generated: {output_path}")

if __name__ == "__main__":
    create_final_master_table()
