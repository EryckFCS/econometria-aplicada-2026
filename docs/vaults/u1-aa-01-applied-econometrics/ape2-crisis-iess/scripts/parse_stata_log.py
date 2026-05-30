#!/usr/bin/env python3
"""Parsear el log de Stata más reciente y generar la tabla de unit root tests.

Ubicación del log por defecto:
 logs/replicate_analysis_ultra_robust.log

Salida:
 reports/unit_root_tests.xlsx
 reports/unit_root_tests.docx
"""
import re
from pathlib import Path
import pandas as pd
from docx import Document


PROJECT_ROOT = Path(__file__).parents[1]
LOG_PATH = PROJECT_ROOT / "logs" / "master_econometric_pipeline.log"
OUT_DIR = PROJECT_ROOT / "reports"
OUT_DIR.mkdir(exist_ok=True)


def normalize_varname(s: str) -> str:
    s = s.lower()
    s = s.replace("~", "")
    # keep alnum only
    return re.sub(r"[^a-z0-9]", "", s)


def stars_from_pvalue(pvalue: float | None) -> str:
    if pvalue is None:
        return ""
    if pvalue < 0.001:
        return "***"
    if pvalue < 0.01:
        return "**"
    if pvalue < 0.05:
        return "*"
    return ""


def stars_from_za(statistic: float | None) -> str:
    if statistic is None:
        return ""
    if statistic <= -5.57:
        return "***"
    if statistic <= -5.08:
        return "**"
    if statistic <= -4.82:
        return "*"
    return ""


def parse_log(text: str):
    # Updated pattern to capture (Levels) or (Diffs) - making type optional for compatibility
    var_block_pattern = re.compile(
        r">>> Variable(?:\s*\((?P<type>Levels|Diffs)\))?:\s*(?P<var>[A-Z]{2}|D\.[A-Z]{2})\n(?P<body>[\s\S]*?)(?=(>>> Variable|\n\. \* 5\.|\Z))"
    )
    adf_pattern = re.compile(r"Augmented Dickey–Fuller test for unit root[\s\S]*?Z\(t\)\s*(?P<zt>-?\d+\.\d+)", re.S)
    pp_pattern = re.compile(r"Phillips–Perron test for unit root[\s\S]*?Z\(t\)\s*(?P<zt>-?\d+\.\d+)", re.S)
    adf_p_pattern = re.compile(r"Augmented Dickey–Fuller test for unit root[\s\S]*?MacKinnon approximate p-value for Z\(t\) = (?P<p>\d+\.\d+)", re.S)
    pp_p_pattern = re.compile(r"Phillips–Perron test for unit root[\s\S]*?MacKinnon approximate p-value for Z\(t\) = (?P<p>\d+\.\d+)", re.S)
    za_pattern = re.compile(r"Minimum t-statistic\s*(?P<tmin>-?\d+\.\d+)\s*at\s*(?P<year>\d{4})")

    parsed = {"Levels": {}, "Diffs": {}}
    for match in var_block_pattern.finditer(text):
        vtype = match.group("type") or "Levels"  # Levels or Diffs
        raw_var = match.group("var")
        # clean variable name (remove D. for diffs)
        var = raw_var.replace("D.", "")
        
        body = match.group("body")
        adf_m = adf_pattern.search(body)
        pp_m = pp_pattern.search(body)
        adf_p_m = adf_p_pattern.search(body)
        pp_p_m = pp_p_pattern.search(body)
        za_m = za_pattern.search(body)

        parsed[vtype][var] = {
            "adf": float(adf_m.group("zt")) if adf_m else None,
            "pp": float(pp_m.group("zt")) if pp_m else None,
            "adf_p": float(adf_p_m.group("p")) if adf_p_m else None,
            "pp_p": float(pp_p_m.group("p")) if pp_p_m else None,
            "za": float(za_m.group("tmin")) if za_m else None,
            "break": int(za_m.group("year")) if za_m else None,
        }

    return parsed


def main():
    if not LOG_PATH.exists():
        print(f"Log no encontrado: {LOG_PATH}")
        return

    text = LOG_PATH.read_text(encoding="utf-8", errors="ignore")

    vars_order = ["AF", "GM", "MI", "HO", "FL"]
    parsed = parse_log(text)

    rows = []
    for short in vars_order:
        # Levels
        l_vals = parsed["Levels"].get(short, {})
        # Diffs
        d_vals = parsed["Diffs"].get(short, {})
        
        # Level Stats
        a_l = l_vals.get("adf")
        p_l = l_vals.get("pp")
        a_l_p = l_vals.get("adf_p")
        p_l_p = l_vals.get("pp_p")
        z_l = l_vals.get("za")
        b_l = l_vals.get("break")

        # Diff Stats
        a_d = d_vals.get("adf")
        p_d = d_vals.get("pp")
        a_d_p = d_vals.get("adf_p")
        p_d_p = d_vals.get("pp_p")
        z_d = d_vals.get("za")
        b_d = d_vals.get("break")

        rows.append(
            {
                "Variable": short,
                "ADF_Level": f"{a_l:.3f}{stars_from_pvalue(a_l_p)}" if a_l is not None else "-",
                "PP_Level": f"{p_l:.3f}{stars_from_pvalue(p_l_p)}" if p_l is not None else "-",
                "ZA_Level": f"{z_l:.3f}{stars_from_za(z_l)}" if z_l is not None else "-",
                "Break_Level": b_l if b_l is not None else "-",
                "ADF_Diff": f"{a_d:.3f}{stars_from_pvalue(a_d_p)}" if a_d is not None else "-",
                "PP_Diff": f"{p_d:.3f}{stars_from_pvalue(p_d_p)}" if p_d is not None else "-",
                "ZA_Diff": f"{z_d:.3f}{stars_from_za(z_d)}" if z_d is not None else "-",
                "Break_Diff": b_d if b_d is not None else "-",
                "I(d)": "I(1)",
            }
        )

    df = pd.DataFrame(rows)

    # reorder columns to match the two-block layout
    df_out = df[
        [
            "Variable",
            "ADF_Level",
            "PP_Level",
            "ZA_Level",
            "Break_Level",
            "ADF_Diff",
            "PP_Diff",
            "ZA_Diff",
            "Break_Diff",
            "I(d)",
        ]
    ]

    df_export = df_out.copy()
    df_export.columns = ["Serie", "ADF", "PP", "ZA", "Break", "ADF", "PP", "ZA", "Break", "I(d)"]

    xlsx_path = OUT_DIR / "unit_root_tests.xlsx"
    df_export.to_excel(xlsx_path, index=False)
    print(f"Generado Excel: {xlsx_path}")

    # Generar Word con una tabla similar
    doc = Document()
    doc.add_paragraph("Table 4")
    doc.add_paragraph("Unit root test results.")

    table = doc.add_table(rows=2 + len(df_out), cols=10)
    table.style = "Table Grid"

    # Header row 1: grouped headings
    row1 = table.rows[0].cells
    row1[0].text = ""
    row1[1].text = "Test-statistics value at Level"
    row1[5].text = "Test-statistics value at first difference"
    row1[9].text = "I(d)"
    row1[1].merge(row1[4])
    row1[5].merge(row1[8])

    # Header row 2: subheadings
    row2 = table.rows[1].cells
    header_titles = [
        "",
        "(ADF)",
        "(PP)",
        "(ZA)",
        "Break",
        "(ADF)",
        "(PP)",
        "(ZA)",
        "Break",
        "",
    ]
    for i, title in enumerate(header_titles):
        row2[i].text = title

    # Data rows
    for ridx, row in enumerate(df_out.itertuples(index=False), start=2):
        cells = table.rows[ridx].cells
        cells[0].text = str(row.Variable)
        cells[1].text = str(row.ADF_Level)
        cells[2].text = str(row.PP_Level)
        cells[3].text = str(row.ZA_Level)
        cells[4].text = str(row.Break_Level)
        cells[5].text = str(row.ADF_Diff)
        cells[6].text = str(row.PP_Diff)
        cells[7].text = str(row.ZA_Diff)
        cells[8].text = str(row.Break_Diff)
        cells[9].text = str(row[9])

    docx_path = OUT_DIR / "unit_root_tests.docx"
    doc.save(docx_path)
    print(f"Generado Word: {docx_path}")


if __name__ == "__main__":
    main()
