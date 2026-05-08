#!/usr/bin/env python3
"""Parsear el log de Stata y generar Excel y Word con la tabla de unit root tests.

Ubicación del log por defecto:
 logs/stata_master_table.log

Salida:
 reports/unit_root_tests.xlsx
 reports/unit_root_tests.docx
"""
import re
from pathlib import Path
import pandas as pd
from docx import Document


PROJECT_ROOT = Path(__file__).parents[1]
LOG_PATH = PROJECT_ROOT / "logs" / "stata_master_table.log"
OUT_DIR = PROJECT_ROOT / "reports"
OUT_DIR.mkdir(exist_ok=True)


def normalize_varname(s: str) -> str:
    s = s.lower()
    s = s.replace("~", "")
    # keep alnum only
    return re.sub(r"[^a-z0-9]", "", s)


def stars_from_pvalue(pvalue: float | None) -> str:
    if pvalue is None:
        return ""
    if pvalue < 0.001:
        return "***"
    if pvalue < 0.01:
        return "**"
    if pvalue < 0.05:
        return "*"
    return ""


def stars_from_za(statistic: float | None) -> str:
    if statistic is None:
        return ""
    if statistic <= -5.57:
        return "***"
    if statistic <= -5.08:
        return "**"
    if statistic <= -4.82:
        return "*"
    return ""


def parse_log(text: str):
    # Patterns for the statistic values; we map them by order because Stata truncates
    # variable names in the console output.
    adf_pattern = re.compile(r"Augmented Dickey–Fuller test for unit root[\s\S]*?Z\(t\)\s*(?P<zt>-?\d+\.\d+)", re.S)
    pp_pattern = re.compile(r"Phillips–Perron test for unit root[\s\S]*?Z\(t\)\s*(?P<zt>-?\d+\.\d+)", re.S)
    adf_p_pattern = re.compile(r"Augmented Dickey–Fuller test for unit root[\s\S]*?MacKinnon approximate p-value for Z\(t\) = (?P<p>\d+\.\d+)", re.S)
    pp_p_pattern = re.compile(r"Phillips–Perron test for unit root[\s\S]*?MacKinnon approximate p-value for Z\(t\) = (?P<p>\d+\.\d+)", re.S)
    za_pattern = re.compile(r"Zivot-Andrews unit root test for[\s\S]*?Minimum t-statistic\s*(?P<tmin>-?\d+\.\d+) at (?P<year>\d{4})", re.S)

    adf_values = [float(m.group("zt")) for m in adf_pattern.finditer(text)]
    pp_values = [float(m.group("zt")) for m in pp_pattern.finditer(text)]
    adf_p_values = [float(m.group("p")) for m in adf_p_pattern.finditer(text)]
    pp_p_values = [float(m.group("p")) for m in pp_p_pattern.finditer(text)]
    za_matches = [(float(m.group("tmin")), int(m.group("year"))) for m in za_pattern.finditer(text)]

    vars_full = [
        ("ln_afiliados_iess", "AF"),
        ("ln_fuerza_laboral", "FL"),
        ("ln_embi_ecuador", "EM"),
        ("ln_gdp_pc_ppp", "GDP"),
        ("ln_sbu", "SBU"),
    ]

    adf = {}
    pp = {}
    adf_p = {}
    pp_p = {}
    za = {}

    for idx, (full, _) in enumerate(vars_full):
        level_pos = idx * 2
        diff_pos = level_pos + 1

        if level_pos < len(adf_values):
            adf[normalize_varname(full)] = adf_values[level_pos]
        if level_pos < len(pp_values):
            pp[normalize_varname(full)] = pp_values[level_pos]
        if idx < len(adf_p_values):
            adf_p[normalize_varname(full)] = adf_p_values[idx]
        if idx < len(pp_p_values):
            pp_p[normalize_varname(full)] = pp_p_values[idx]
        if level_pos < len(za_matches):
            tmin, year = za_matches[level_pos]
            za[normalize_varname(full)] = {"tmin": tmin, "year": year}

    for idx, (full, _) in enumerate(vars_full):
        dkey = normalize_varname("d." + full)
        diff_pos = idx * 2 + 1
        if diff_pos < len(adf_values):
            adf[dkey] = adf_values[diff_pos]
        if diff_pos < len(pp_values):
            pp[dkey] = pp_values[diff_pos]
        if idx < len(adf_p_values):
            adf_p[dkey] = adf_p_values[idx]
        if idx < len(pp_p_values):
            pp_p[dkey] = pp_p_values[idx]
        if diff_pos < len(za_matches):
            tmin, year = za_matches[diff_pos]
            za[dkey] = {"tmin": tmin, "year": year}

    return adf, pp, adf_p, pp_p, za


def main():
    if not LOG_PATH.exists():
        print(f"Log no encontrado: {LOG_PATH}")
        return

    text = LOG_PATH.read_text(encoding="utf-8", errors="ignore")

    # variables objetivo en el orden esperado (según el do-file)
    vars_full = [
        ("ln_afiliados_iess", "AF"),
        ("ln_fuerza_laboral", "FL"),
        ("ln_embi_ecuador", "EM"),
        ("ln_gdp_pc_ppp", "GDP"),
        ("ln_sbu", "SBU"),
    ]

    adf, pp, adf_p, pp_p, za = parse_log(text)

    rows = []
    for full, short in vars_full:
        key = normalize_varname(full)

        # level values
        a_l = adf.get(key)
        p_l = pp.get(key)
        a_l_p = adf_p.get(key)
        p_l_p = pp_p.get(key)
        z_l = za.get(key, {}).get("tmin") if za.get(key) else None
        by = za.get(key, {}).get("year") if za.get(key) else None

        # difference keys: in log the diff variable often appears as D.<var>
        dkey = normalize_varname("d." + full)
        a_d = adf.get(dkey)
        p_d = pp.get(dkey)
        a_d_p = adf_p.get(dkey)
        p_d_p = pp_p.get(dkey)
        z_d = za.get(dkey, {}).get("tmin") if za.get(dkey) else None
        bdy = za.get(dkey, {}).get("year") if za.get(dkey) else None

        rows.append(
            {
                "Variable": short,
                "ADF_Level": f"{a_l:.3f}{stars_from_pvalue(a_l_p)}" if a_l is not None else "",
                "PP_Level": f"{p_l:.3f}{stars_from_pvalue(p_l_p)}" if p_l is not None else "",
                "ZA_Level": f"{z_l:.3f}{stars_from_za(z_l)}" if z_l is not None else "",
                "Break_Level": by if by is not None else "",
                "ADF_Diff": f"{a_d:.3f}{stars_from_pvalue(a_d_p)}" if a_d is not None else "",
                "PP_Diff": f"{p_d:.3f}{stars_from_pvalue(p_d_p)}" if p_d is not None else "",
                "ZA_Diff": f"{z_d:.3f}{stars_from_za(z_d)}" if z_d is not None else "",
                "Break_Diff": bdy if bdy is not None else "",
                "I(d)": "I(1)",
            }
        )

    df = pd.DataFrame(rows)

    # reorder columns to match the two-block layout
    df_out = df[
        [
            "Variable",
            "ADF_Level",
            "PP_Level",
            "ZA_Level",
            "Break_Level",
            "ADF_Diff",
            "PP_Diff",
            "ZA_Diff",
            "Break_Diff",
            "I(d)",
        ]
    ]

    df_export = df_out.copy()
    df_export.columns = ["", "(ADF)", "(PP)", "(ZA)", "Break", "(ADF)", "(PP)", "(ZA)", "Break", ""]

    xlsx_path = OUT_DIR / "unit_root_tests.xlsx"
    df_export.to_excel(xlsx_path, index=False)
    print(f"Generado Excel: {xlsx_path}")

    # Generar Word con una tabla similar
    doc = Document()
    doc.add_paragraph("Table 4")
    doc.add_paragraph("Unit root test results.")

    table = doc.add_table(rows=2 + len(df_out), cols=10)
    table.style = "Table Grid"

    # Header row 1: grouped headings
    row1 = table.rows[0].cells
    row1[0].text = ""
    row1[1].text = "Test-statistics value at Level"
    row1[5].text = "Test-statistics value at first difference"
    row1[9].text = "I(d)"
    row1[1].merge(row1[4])
    row1[5].merge(row1[8])

    # Header row 2: subheadings
    row2 = table.rows[1].cells
    header_titles = [
        "",
        "(ADF)",
        "(PP)",
        "(ZA)",
        "Break",
        "(ADF)",
        "(PP)",
        "(ZA)",
        "Break",
        "",
    ]
    for i, title in enumerate(header_titles):
        row2[i].text = title

    # Data rows
    for ridx, row in enumerate(df_out.itertuples(index=False), start=2):
        cells = table.rows[ridx].cells
        cells[0].text = str(row.Variable)
        cells[1].text = str(row.ADF_Level)
        cells[2].text = str(row.PP_Level)
        cells[3].text = str(row.ZA_Level)
        cells[4].text = str(row.Break_Level)
        cells[5].text = str(row.ADF_Diff)
        cells[6].text = str(row.PP_Diff)
        cells[7].text = str(row.ZA_Diff)
        cells[8].text = str(row.Break_Diff)
        cells[9].text = str(row[9])

    docx_path = OUT_DIR / "unit_root_tests.docx"
    doc.save(docx_path)
    print(f"Generado Word: {docx_path}")


if __name__ == "__main__":
    main()
